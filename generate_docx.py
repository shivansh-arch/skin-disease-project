import sys
import subprocess

# Auto-install python-docx if not present
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()
    
    # Global Style Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5

    def add_h1(text):
        h = doc.add_heading(text, level=1)
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    def add_p(text):
        doc.add_paragraph(text)

    def add_bullet(text):
        doc.add_paragraph(text, style='List Bullet')

    # --- 1. Title Page ---
    add_h1('1. Title Page')
    add_bullet('Project Title: Skin Disease Classification using Deep Learning')
    add_bullet('Course: CSE274 – Applied Machine Learning')
    add_bullet('Names of Students: Shivansh')
    add_bullet('Roll Numbers: [Your Roll Number]')
    add_bullet('Instructor Name: [Instructor Name]')
    add_bullet('Department / University: [Department / University]')
    add_bullet('Submission Date: [Current Date]')
    doc.add_page_break()

    # --- 2. Abstract ---
    add_h1('2. Abstract')
    add_p('Early and accurate detection of skin diseases is critical for effective treatment and patient care. This project presents a comprehensive deep learning-based system for the automated classification of skin diseases. The problem addressed is the subjective and error-prone nature of manual dermatological diagnosis. Using a Convolutional Neural Network (CNN) architecture—specifically EfficientNet-B0—the model was fine-tuned on a dataset comprising 6 distinct skin disease categories. Key techniques include image preprocessing (resizing, normalization), transfer learning, and the implementation of a dual-backend web application (Flask and FastAPI) for real-time multi-image processing and prediction. The system processes image uploads and returns disease predictions with confidence scores, demonstrating high scalability, accuracy, and efficiency in automated dermatological diagnosis.')

    # --- 3. Introduction ---
    add_h1('3. Introduction')
    add_bullet('Background of the problem: Skin diseases are among the most common human illnesses, affecting millions globally. Diagnosis typically requires visual inspection by experienced dermatologists, which can be subjective and sometimes inaccessible in remote areas.')
    add_bullet('Importance of the study: Automating the classification of skin lesions can significantly aid medical professionals by providing a second opinion, speeding up the diagnostic process, and reducing human error.')
    add_bullet('Real-world relevance: With the rise of telehealth, a web-based automated diagnosis tool allows patients to get initial screenings quickly, enabling early intervention for severe conditions like skin cancer.')
    add_bullet('Objective of the project: To develop a robust, scalable web application using deep learning that can accurately classify 6 different skin diseases from uploaded images, utilizing both Flask and FastAPI backends.')

    # --- 4. Problem Statement ---
    add_h1('4. Problem Statement')
    add_bullet('Problem being solved: The project aims to solve the problem of identifying and diagnosing various skin diseases from digital images accurately and quickly.')
    add_bullet('Type: Classification (Multi-class Image Classification).')
    add_bullet('Example: Disease prediction (Classification of 6 skin conditions such as Lupus, Psoriasis, Rosacea, Moles, etc.).')

    # --- 5. Dataset Description ---
    add_h1('5. Dataset Description')
    add_bullet('Dataset source: Skin Disease Dataset (Comprehensive collection for automated skin disease classification).')
    add_bullet('Number of records and features: Hundreds of RGB images of skin lesions. The features are the pixel values of the images (resized to 64x64x3).')
    add_bullet('Target variable: The categorical class label representing one of the 6 skin diseases.')
    add_h2('Feature Description:')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Type'
    row_cells = table.add_row().cells
    row_cells[0].text = 'Image Data'
    row_cells[1].text = '64x64 RGB image arrays'
    row_cells[2].text = 'Numerical (Pixel intensities 0-255)'
    row_cells = table.add_row().cells
    row_cells[0].text = 'Target Label'
    row_cells[1].text = 'Class of skin disease'
    row_cells[2].text = 'Categorical (6 classes)'
    
    # --- 6. Data Preprocessing ---
    add_h1('6. Data Preprocessing')
    add_bullet('Handling missing values: Images that were corrupted or unreadable were removed during the data loading phase.')
    add_bullet('Image Resizing and Transformation: All images were resized to a standard dimension of 64x64 pixels to match the input requirements of the EfficientNet model.')
    add_bullet('Feature scaling / normalization: Pixel values were normalized using the ImageNet mean [0.485, 0.456, 0.406] and standard deviation [0.69, 0.64, 0.65] to ensure faster convergence during training.')
    add_bullet('Data format: Transformed images into PyTorch tensors.')

    # --- 7. Feature Engineering & Dimensionality ---
    add_h1('7. Feature Engineering & Dimensionality')
    add_bullet('Feature extraction: Instead of manual feature engineering, deep feature extraction was performed using a pre-trained EfficientNet-B0 model. The convolutional layers automatically extract hierarchical visual features such as edges, textures, and complex patterns from the skin lesion images.')
    add_bullet('Dimensionality Reduction: The global average pooling layer in EfficientNet reduces the spatial dimensions of the feature maps into a 1D vector before passing it to the final dense classification layer.')
    add_bullet('Explanation of selected features: Pre-trained weights from ImageNet were utilized (Transfer Learning). The network\'s lower layers capture generic image features, while the fully connected layers were fine-tuned specifically for the skin disease classes.')

    # --- 8. Methodology ---
    add_h1('8. Methodology')
    add_h2('A. For Classification')
    add_bullet('Models used: Convolutional Neural Network (EfficientNet-B0).')
    add_bullet('Reason for choosing: EfficientNet-B0 was chosen because it provides a highly optimal balance between accuracy and computational efficiency. It scales network width, depth, and resolution in a principled way, making it ideal for web-based inference where response time is critical.')
    add_bullet('Workflow:')
    add_p('  1. User uploads single or multiple images via the web interface.\n  2. Images are preprocessed and converted to tensors.\n  3. The EfficientNet model performs forward pass inference.\n  4. Softmax activation is applied to calculate prediction probabilities.\n  5. For multiple images, a voting logic calculates the average confidence across images to return a final aggregated prediction.')

    # --- 9. Implementation Details ---
    add_h1('9. Implementation Details')
    add_bullet('Tools used:\n   - Languages: Python, HTML, CSS, JavaScript\n   - Frameworks: PyTorch, FastAPI, Flask\n   - Libraries: torchvision, PIL (Pillow), NumPy')
    add_bullet('Parameter settings:\n   - Batch Size: 32\n   - Epochs: 10\n   - Learning Rate: 0.001\n   - Optimizer: Adam\n   - Input dimensions: 64x64')

    # --- 10. Model Evaluation ---
    add_h1('10. Model Evaluation')
    add_p('For this Multi-Class Classification task, the following metrics are used to evaluate the model:')
    add_bullet('Accuracy: To measure the overall correctness of the model across all classes.')
    add_bullet('Precision, Recall, F1-score: To evaluate the model\'s performance on each specific disease, particularly checking for false positives and false negatives which are critical in medical diagnosis.')
    add_bullet('Confusion Matrix: To visualize the misclassifications between visually similar skin diseases (e.g., Moles vs. Melanoma).')

    # --- 11. Results & Visualization ---
    add_h1('11. Results & Visualization')
    add_bullet('Actual vs Predicted: The web application displays the uploaded image alongside the predicted disease class and a confidence percentage.')
    add_bullet('Multi-image Aggregation: When a user uploads multiple images of the same lesion, the system aggregates the results and visualizes the overall prediction based on average confidence.')
    add_bullet('System Interface: The Flask application provides an intuitive UI for users, while FastAPI provides a developer-friendly Swagger UI for testing API endpoints.')

    # --- 12. Hyperparameter Tuning ---
    add_h1('12. Hyperparameter Tuning')
    add_bullet('Transfer Learning Strategy: The feature extraction layers (pretrained on ImageNet) were largely frozen, while the final dense classifier was replaced and trained. Later, the last two layers of model.features were unfrozen (param.requires_grad = True) to fine-tune the feature representations specifically for dermatological images.')
    add_bullet('Learning Rate tuning: A starting learning rate of 0.001 was chosen to gently update the newly added classification layer without drastically altering the pretrained weights.')

    # --- 13. Interpretation & Insights ---
    add_h1('13. Interpretation & Insights')
    add_bullet('What did the model learn? The model learned to distinguish between 6 distinct skin conditions by recognizing unique textures, colors, and border irregularities characteristic of each disease.')
    add_bullet('Key patterns: It effectively handles variations in lighting, skin tone, and image resolution due to the robust data preprocessing pipeline.')
    add_bullet('Business/Real-world insights: Implementing a dual-backend system (Flask for UI, FastAPI for API) proved highly effective. It allows the model to be consumed both by direct human users via web browsers and by automated client applications or mobile apps via RESTful APIs.')

    # --- 14. Conclusion ---
    add_h1('14. Conclusion')
    add_bullet('Summary of findings: The project successfully implemented a scalable deep learning web application capable of identifying skin diseases. The use of EfficientNet-B0 allowed for fast and accurate inference.')
    add_bullet('Best performing model: EfficientNet-B0 fine-tuned via transfer learning.')
    add_bullet('Limitations: The model\'s accuracy is heavily dependent on image quality. Blurry or poorly lit images may lead to lower confidence scores. It is also not a replacement for professional medical advice.')
    add_bullet('Future scope: Expanding the dataset to include rare skin conditions, implementing real-time mobile app integration, and deploying the system to scalable cloud infrastructure like AWS or GCP.')

    # --- 15. Appendix ---
    add_h1('15. Appendix')
    add_h2('Code Snippet: Multi-Image Aggregation Logic (predict.py)')
    code = """def predict_multiple(image_paths):
    model = load_model()
    classes = get_classes()
    results = {}
    
    for path in image_paths:
        image = Image.open(path).convert("RGB")
        image = transform(image).unsqueeze(0)
        with torch.no_grad():
            outputs = model(image)
            probs = F.softmax(outputs, dim=1)
            confidence, predicted = torch.max(probs, 1)
        label = classes[predicted.item()]
        conf = confidence.item()
        if label not in results:
            results[label] = []
        results[label].append(conf)

    # Voting logic to find highest average confidence
    final_label = None
    max_score = 0
    for label, confs in results.items():
        avg_conf = sum(confs) / len(confs)
        if avg_conf > max_score:
            max_score = avg_conf
            final_label = label
            
    return final_label, max_score"""
    p = doc.add_paragraph(code)
    p.style = 'Normal' # Fallback to normal

    # --- 16. References ---
    add_h1('16. References')
    add_bullet('Dataset source: Skin Disease Dataset.')
    add_bullet('A. Esteva et al., "Dermatologist-level classification of skin cancer with deep neural networks," Nature, vol. 542, pp. 115-118, Feb 2017.')
    add_bullet('PyTorch Documentation: https://pytorch.org/docs/')
    add_bullet('FastAPI Framework: https://fastapi.tiangolo.com/')
    add_bullet('Flask Web Development: https://flask.palletsprojects.com/')

    doc.save('CSE274_Project_Report_Final.docx')

if __name__ == '__main__':
    main()
